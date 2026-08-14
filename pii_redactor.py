import os
import re
import docx
from faker import Faker

class Redactor:
    def __init__(self, lang="en_US"):
        self.fake = Faker(lang)
        self.cache = {}
        self._build_patterns()

    def _build_patterns(self):
        self.patterns = [
            ("EMAIL_ADDRESS", re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")),
            ("PHONE_NUMBER", re.compile(r"(\+?\d{1,3}[\s-]?)?\(?\d{3,5}\)?[\s-]?\d{3,5}[\s-]?\d{3,4}")),
            ("DATE_OF_BIRTH", re.compile(r"\b(0[1-9]|[12][0-9]|3[01])[-/.](0[1-9]|1[012])[-/.](19|20)\d\d\b")),
            ("US_SSN", re.compile(r"\b\d{3}-\d{2}-\d{4}\b")),
            ("CREDIT_CARD", re.compile(r"\b(?:\d[ -]*?){13,16}\b")),
            ("IP_ADDRESS", re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b")),
        ]

    def _get_fake(self, tag: str, text: str) -> str:
        if text in self.cache:
            return self.cache[text]

        if "EMAIL" in tag:
            val = self.fake.email()
        elif "PHONE" in tag:
            val = self.fake.phone_number()
        elif "SSN" in tag:
            val = self.fake.ssn()
        elif "CARD" in tag:
            val = self.fake.credit_card_number()
        elif "DATE" in tag:
            val = self.fake.date()
        elif "IP" in tag:
            val = self.fake.ipv4()
        else:
            val = f"[REDACTED_{tag}]"

        self.cache[text] = val
        return val

    def clean_text(self, text: str) -> str:
        if not text or not text.strip():
            return text

        res = text
        for tag, pattern in self.patterns:
            matches = list(pattern.finditer(res))
            for match in reversed(matches):
                old_val = match.group()
                new_val = self._get_fake(tag, old_val)
                start, end = match.span()
                res = res[:start] + new_val + res[end:]

        return res

    def run(self, in_file: str, out_file: str):
        if not os.path.exists(in_file):
            print(f"Error: File '{in_file}' not found.")
            return

        print(f"Opening '{in_file}'...")
        doc = docx.Document(in_file)

        print("Processing paragraphs...")
        for p in doc.paragraphs:
            if p.text.strip():
                p.text = self.clean_text(p.text)

        print("Processing tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        cell.text = self.clean_text(cell.text)

        doc.save(out_file)
        print(f"SUCCESS! Redacted file saved as: {out_file}")


if __name__ == "__main__":
    print("Script started...")
    app = Redactor()
    app.run("Assignment_Document.docx", "Redacted_Document.docx")